#!/usr/bin/env python3
"""生成Python基础期末考试试卷（ABC三卷）"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import json
import os

def set_run_font(run, font_name='宋体', font_size=12, bold=False):
    """设置字体"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_heading(doc, text, level=1, center=True):
    """添加标题"""
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    size = 18 if level == 1 else (14 if level == 2 else 12)
    set_run_font(run, '黑体', size, bold=True)
    return p

def add_paragraph(doc, text, font_name='宋体', font_size=12, bold=False, indent=False, spacing_before=0, spacing_after=6):
    """添加段落"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(spacing_before)
    p.paragraph_format.space_after = Pt(spacing_after)
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run(text)
    set_run_font(run, font_name, font_size, bold)
    return p

def add_bullet_point(doc, text, font_name='宋体', font_size=12):
    """添加项目符号段落"""
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    set_run_font(run, font_name, font_size)
    return p

def generate_exam_doc(exam_data, output_path):
    """生成一份试卷文档"""
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)

    exam_info = exam_data['examInstructions']
    scoring = exam_data['scoringCriteria']

    # ============ 封面信息 ============
    add_heading(doc, exam_data['title'], level=1)
    add_heading(doc, '（机考实操试卷适用）', level=2)

    # ============ 一、考试须知 ============
    add_heading(doc, '一、考试须知', level=2, center=False)
    for item in exam_info.get('考试须知', []):
        # 替换时长为120分钟
        item = item.replace('90分钟', '120分钟')
        add_bullet_point(doc, item)

    # ============ 二、答题要求 ============
    add_heading(doc, '二、答题要求', level=2, center=False)
    for item in exam_info.get('答题要求', []):
        add_bullet_point(doc, item)

    # ============ 三、考试纪律 ============
    add_heading(doc, '三、考试纪律', level=2, center=False)
    for item in exam_info.get('考试纪律', []):
        add_bullet_point(doc, item)

    # ============ 四、交卷说明 ============
    add_heading(doc, '四、交卷说明', level=2, center=False)
    for item in exam_info.get('交卷说明', []):
        add_bullet_point(doc, item)

    # ============ 五、评分标准 ============
    add_heading(doc, '五、评分标准', level=2, center=False)
    add_paragraph(doc, f"总分：{scoring['总分']}分  |  及格分数：{scoring['及格分']}分", bold=True)

    # 分值分布表
    add_paragraph(doc, '（一）函数定义与基础操作（第一题至第八题）', bold=True, spacing_before=12)

    table1 = doc.add_table(rows=1, cols=3)
    table1.style = 'Table Grid'
    headers = ['题号', '题目名称', '分值']
    for i, header in enumerate(headers):
        cell = table1.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            for run in p.runs:
                set_run_font(run, '黑体', 10, bold=True)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i in range(1, 9):
        q_key = f'q{i}'
        q_info = scoring['分值分布']['函数定义与基础操作']['题目分值'].get(q_key)
        if q_info:
            row = table1.add_row().cells
            row[0].text = f'第{i}题'
            row[1].text = q_info['名称']
            row[2].text = str(q_info['分值'])
            for j, cell in enumerate(row):
                for p in cell.paragraphs:
                    for run in p.runs:
                        set_run_font(run, '宋体', 9)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_paragraph(doc, '', spacing_before=12)
    add_paragraph(doc, '（二）输出结果与综合应用（第九题至第十二题）', bold=True)

    table2 = doc.add_table(rows=1, cols=3)
    table2.style = 'Table Grid'
    for i, header in enumerate(headers):
        cell = table2.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            for run in p.runs:
                set_run_font(run, '黑体', 10, bold=True)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i in range(9, 13):
        q_key = f'q{i}'
        q_info = scoring['分值分布']['输出结果与综合应用']['题目分值'].get(q_key)
        if q_info:
            row = table2.add_row().cells
            row[0].text = f'第{i}题'
            row[1].text = q_info['名称']
            row[2].text = str(q_info['分值'])
            for j, cell in enumerate(row):
                for p in cell.paragraphs:
                    for run in p.runs:
                        set_run_font(run, '宋体', 9)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_paragraph(doc, '', spacing_before=12)
    add_paragraph(doc, '评分细则：', bold=True)
    for key, value in scoring['评分细则'].items():
        add_bullet_point(doc, f"{value['占比']} - {value['说明']}")

    # ============ 六、题目内容 ============
    add_heading(doc, '六、题目内容', level=2, center=False)

    for i, q in enumerate(exam_data['questions'], 1):
        # 题目标题
        add_paragraph(doc, f"第{i}题 ({q['scoring']['points']}分)", bold=True, spacing_before=12)
        add_paragraph(doc, f"【{q['type'].upper()}】{q['title']}", bold=True)

        # 题目描述（完整instruction）
        instruction_lines = q['instruction'].split('\n')
        for line in instruction_lines:
            line = line.strip()
            if line:
                add_paragraph(doc, line, font_size=11, indent=True)

        # 初始代码
        add_paragraph(doc, f"初始代码：", bold=True)
        code_p = doc.add_paragraph()
        code_p.paragraph_format.left_indent = Inches(0.5)
        code_run = code_p.add_run(q['initialCode'].strip())
        set_run_font(code_run, 'Consolas', 10)

        # 预期输出
        expected = q['testConfig'].get('expected', '')
        if expected:
            add_paragraph(doc, f"预期输出：", bold=True)
            exp_p = doc.add_paragraph()
            exp_p.paragraph_format.left_indent = Inches(0.5)
            exp_run = exp_p.add_run(str(expected))
            set_run_font(exp_run, 'Consolas', 10)

        # 评分标准（每道题后）
        add_paragraph(doc, f"评分标准：{q['scoring']['criteria']}", bold=True, font_name='黑体')

        add_paragraph(doc, '')

    # ============ 底部说明 ============
    add_paragraph(doc, '')
    add_paragraph(doc, '【说明】本试卷共12题，总分100分，考试时长120分钟。', bold=False)
    add_paragraph(doc, '考试地址：https://py-learning.pages.dev/', bold=False)

    # 保存文档
    doc.save(output_path)
    print(f"已生成：{output_path}")

def main():
    # 读取JSON文件
    exams = []
    for version in ['A', 'B', 'C']:
        json_path = f'/workspace/py-learning/public/data/exam/final_exam_{version}.json'
        with open(json_path, 'r', encoding='utf-8') as f:
            exams.append((version, json.load(f)))

    # 输出目录
    output_dir = '/workspace/py-learning/public/data/exam/docs'
    os.makedirs(output_dir, exist_ok=True)

    # 生成三份试卷
    for version, exam in exams:
        output_path = os.path.join(output_dir, f"期末考试_{version}卷.docx")
        generate_exam_doc(exam, output_path)

    print(f"\n所有试卷已生成到：{output_dir}")

if __name__ == '__main__':
    main()
